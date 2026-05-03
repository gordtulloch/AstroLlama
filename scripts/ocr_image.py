"""
OCR script — converts JPG images in a folder to text or DOCX files using Tesseract.

Usage:
    python scripts/ocr_image.py --input path/to/images
    python scripts/ocr_image.py --input path/to/images --combine
    python scripts/ocr_image.py --input path/to/images --combine --docx
    python scripts/ocr_image.py --input path/to/images --rotate 90
    python scripts/ocr_image.py --input path/to/images --save-preprocessed
    python scripts/ocr_image.py --input path/to/images --output path/to/output
    python scripts/ocr_image.py --input path/to/images --dpi 300
"""
from __future__ import annotations

import argparse
import logging
import sys
from pathlib import Path

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)-8s %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger(__name__)

_DEFAULT_INPUT = Path.cwd()
_SUPPORTED = {".jpg", ".jpeg"}


def _configure_tesseract() -> None:
    """Point pytesseract at the default Windows Tesseract install if needed."""
    import pytesseract

    if sys.platform == "win32":
        win_default = Path(r"C:\Program Files\Tesseract-OCR\tesseract.exe")
        if win_default.exists():
            pytesseract.pytesseract.tesseract_cmd = str(win_default)


def _auto_rotate(img: "Image.Image") -> "Image.Image":
    """Use Tesseract OSD to detect orientation and rotate the image upright."""
    import pytesseract

    try:
        osd = pytesseract.image_to_osd(img, output_type=pytesseract.Output.DICT)
        angle = osd.get("rotate", 0)
        if angle:
            logger.info("  Auto-rotating image by %d°", angle)
            img = img.rotate(angle, expand=True)
    except Exception as exc:
        logger.debug("OSD rotation detection failed (skipping): %s", exc)
    return img


def _preprocess(img: "Image.Image") -> "Image.Image":
    """Convert to greyscale and enhance contrast/sharpness for better OCR accuracy."""
    from PIL import ImageEnhance, ImageFilter

    img = img.convert("L")  # greyscale
    img = ImageEnhance.Contrast(img).enhance(2.0)
    img = ImageEnhance.Sharpness(img).enhance(2.0)
    img = img.filter(ImageFilter.MedianFilter(size=1))  # light denoise
    return img


def _ocr_image_file(
    img_path: Path,
    dpi: int,
    auto_rotate: bool = True,
    force_rotate: int = 0,
    psm: int = 3,
    save_preprocessed: bool = False,
    output_dir: Path | None = None,
) -> str:
    """Open an image, preprocess, optionally rotate, then return OCR'd text."""
    import pytesseract
    from PIL import Image

    img = Image.open(img_path)

    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")

    if dpi:
        img_dpi = img.info.get("dpi", (72, 72))
        scale = dpi / img_dpi[0] if img_dpi[0] > 0 else 1.0
        if scale > 1.01:
            new_size = (int(img.width * scale), int(img.height * scale))
            img = img.resize(new_size, Image.LANCZOS)

    # Manual rotation takes priority over auto-rotate
    if force_rotate:
        img = img.rotate(force_rotate, expand=True)
        logger.info("  Rotated %d°", force_rotate)
    elif auto_rotate:
        img = _auto_rotate(img)

    img = _preprocess(img)

    if save_preprocessed and output_dir:
        pre_path = output_dir / ("pre_" + img_path.stem + ".png")
        img.save(str(pre_path))
        logger.info("  Saved preprocessed image → %s", pre_path.name)

    config = f"--psm {psm}"
    return pytesseract.image_to_string(img, config=config)


def _write_docx(path: Path, sections: list[tuple[str, str]]) -> None:
    """Write a DOCX file from a list of (heading, body_text) tuples."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    for heading, text in sections:
        if heading:
            doc.add_heading(heading, level=1)
        for line in text.splitlines():
            para = doc.add_paragraph(line)
            para.style.font.size = Pt(11)
    doc.save(str(path))


def process_folder(
    input_dir: Path,
    output_dir: Path,
    dpi: int,
    combine: bool = False,
    combined_name: str = "combined",
    auto_rotate: bool = True,
    force_rotate: int = 0,
    psm: int = 3,
    docx: bool = False,
    save_preprocessed: bool = False,
) -> int:
    """OCR all JPG images in input_dir.

    When combine=False writes one file per image; when combine=True writes a single file.
    Output format is .docx when docx=True, otherwise .txt.
    """
    _configure_tesseract()

    ocr_kwargs: dict = dict(
        dpi=dpi,
        auto_rotate=auto_rotate,
        force_rotate=force_rotate,
        psm=psm,
        save_preprocessed=save_preprocessed,
        output_dir=output_dir,
    )

    images = [
        p for p in sorted(input_dir.iterdir())
        if p.is_file() and p.suffix.lower() in _SUPPORTED
    ]

    if not images:
        logger.warning("No JPG images found in: %s", input_dir)
        return 0

    output_dir.mkdir(parents=True, exist_ok=True)
    count = 0

    if combine:
        ext = ".docx" if docx else ".txt"
        stem = Path(combined_name).stem
        combined_path = output_dir / (stem + ext)
        if docx:
            sections: list[tuple[str, str]] = []
            for img_path in images:
                try:
                    text = _ocr_image_file(img_path, **ocr_kwargs)
                    sections.append((img_path.name, text))
                    logger.info("%-40s → %s", img_path.name, combined_path.name)
                    count += 1
                except Exception as exc:
                    logger.error("Failed to process %s: %s", img_path.name, exc)
            _write_docx(combined_path, sections)
        else:
            with combined_path.open("w", encoding="utf-8") as fh:
                for img_path in images:
                    try:
                        text = _ocr_image_file(img_path, **ocr_kwargs)
                        fh.write(f"{'='*80}\n")
                        fh.write(f"SOURCE: {img_path.name}\n")
                        fh.write(f"{'='*80}\n")
                        fh.write(text)
                        fh.write("\n\n")
                        logger.info("%-40s → %s", img_path.name, combined_path.name)
                        count += 1
                    except Exception as exc:
                        logger.error("Failed to process %s: %s", img_path.name, exc)
        logger.info("Combined output written to: %s", combined_path)
    else:
        for img_path in images:
            try:
                text = _ocr_image_file(img_path, **ocr_kwargs)
                if docx:
                    out_path = output_dir / (img_path.stem + ".docx")
                    _write_docx(out_path, [(img_path.name, text)])
                else:
                    out_path = output_dir / (img_path.stem + ".txt")
                    out_path.write_text(text, encoding="utf-8")
                logger.info("%-40s → %s", img_path.name, out_path.name)
                count += 1
            except Exception as exc:
                logger.error("Failed to process %s: %s", img_path.name, exc)

    return count


def main() -> None:
    parser = argparse.ArgumentParser(
        description="OCR JPG images in a folder and write text or DOCX output"
    )
    parser.add_argument(
        "--input",
        type=Path,
        required=True,
        help="Folder containing JPG images",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=None,
        help="Folder for output files (default: current working directory)",
    )
    parser.add_argument(
        "--dpi",
        type=int,
        default=300,
        help="Target DPI for upscaling low-resolution images before OCR (default: 300). "
             "Use 0 to skip upscaling.",
    )
    parser.add_argument(
        "--combine",
        action="store_true",
        help="Combine all OCR'd text into a single file instead of one file per image.",
    )
    parser.add_argument(
        "--combined-name",
        default="combined",
        help="Base filename for the combined output (default: combined). "
             "Extension is added automatically. Only used with --combine.",
    )
    parser.add_argument(
        "--docx",
        action="store_true",
        help="Write output as DOCX instead of plain text.",
    )
    parser.add_argument(
        "--no-rotate",
        action="store_true",
        help="Disable automatic orientation detection and rotation (enabled by default).",
    )
    parser.add_argument(
        "--rotate",
        type=int,
        default=0,
        choices=[0, 90, 180, 270],
        help="Force-rotate all images by this many degrees counter-clockwise before OCR "
             "(overrides auto-rotate). Use 90 or 270 for landscape scans.",
    )
    parser.add_argument(
        "--psm",
        type=int,
        default=3,
        help="Tesseract page segmentation mode (default: 3 = fully automatic). "
             "Try 6 (single uniform block) or 4 (single column) for structured documents.",
    )
    parser.add_argument(
        "--save-preprocessed",
        action="store_true",
        help="Save the preprocessed (greyscale/enhanced) image to the output folder "
             "as pre_<name>.png so you can inspect what Tesseract sees.",
    )
    args = parser.parse_args()

    input_dir: Path = args.input.resolve()
    output_dir: Path = (args.output or Path.cwd()).resolve()

    if not input_dir.exists():
        logger.error("Input folder does not exist: %s", input_dir)
        sys.exit(1)

    logger.info("Input  : %s", input_dir)
    logger.info("Output : %s", output_dir)

    count = process_folder(
        input_dir, output_dir, args.dpi,
        combine=args.combine,
        combined_name=args.combined_name,
        auto_rotate=not args.no_rotate,
        force_rotate=args.rotate,
        psm=args.psm,
        docx=args.docx,
        save_preprocessed=args.save_preprocessed,
    )
    logger.info("Done — %d image(s) converted to text", count)


if __name__ == "__main__":
    main()
